from docx import Document
import pandas as pd
import os

def generate_word_report(job_records, output_file):
    doc = Document()
    doc.add_heading("Hadoop Shell Job Execution Report", level=1)

    for record in job_records:
        doc.add_heading(record['script'], level=2)
        doc.add_paragraph(f"Start Time: {record['start_time']}")
        doc.add_paragraph(f"End Time: {record['end_time']}")
        doc.add_paragraph(f"App ID: {record['app_id']}")
        doc.add_paragraph(f"Status: {record['status']}")
        doc.add_paragraph(f"Command: {record['command']}")
        doc.add_paragraph()

    doc.save(output_file)

def generate_csv_report(job_records, output_file):
    df = pd.DataFrame(job_records)
    df.to_csv(output_file, index=False)
