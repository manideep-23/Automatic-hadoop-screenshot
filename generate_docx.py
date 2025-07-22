import os
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches

# Base URL for ApplicationMaster
APP_MASTER_BASE_URL = "http://resourcemanager-host:8088/cluster/app/"

# Screenshot folder
SCREENSHOT_DIR = "screenshots"
os.makedirs(SCREENSHOT_DIR, exist_ok=True)

def setup_driver():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--window-size=1280,720")
    driver = webdriver.Chrome(ChromeDriverManager().install(), options=chrome_options)
    return driver

def process_file(input_file, output_docx):
    doc = Document()
    driver = setup_driver()

    with open(input_file, "r") as file:
        for idx, line in enumerate(file):
            line = line.strip()
            if not line or line.startswith("#"):
                continue  # skip blank or commented lines

            try:
                command, app_id, status = map(str.strip, line.split(",", 2))
            except ValueError:
                print(f"Skipping invalid line {idx+1}: {line}")
                continue

            doc.add_paragraph(f"Command Executed : {command}")
            doc.add_paragraph(f"Status           : {status}")

            url = f"{APP_MASTER_BASE_URL}{app_id}"
            print(f"[{idx+1}] Opening: {url}")
            try:
                driver.get(url)
                time.sleep(2)  # wait for page to load
                screenshot_path = os.path.join(SCREENSHOT_DIR, f"{app_id}.png")
                driver.save_screenshot(screenshot_path)
                doc.add_picture(screenshot_path, width=Inches(6.0))
            except Exception as e:
                print(f"Error capturing {app_id}: {e}")
                doc.add_paragraph("⚠️ Failed to capture screenshot.")

            doc.add_paragraph("\n" + "-" * 60 + "\n")

    driver.quit()
    doc.save(output_docx)
    print(f"✅ Document generated: {output_docx}")

if __name__ == "__main__":
    process_file("file1.txt", "successful_jobs.docx")
    process_file("file2.txt", "failed_jobs.docx")
