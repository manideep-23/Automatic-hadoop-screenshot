import os
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Inches


def generate_doc(input_file, output_doc, chrome_driver_path, rm_base_url):
    doc = Document()
    screenshots_dir = "screenshots"
    os.makedirs(screenshots_dir, exist_ok=True)

    # Setup Selenium headless Chrome
    options = Options()
    options.add_argument("--headless")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=1200x800")

    driver = webdriver.Chrome(executable_path=chrome_driver_path, options=options)

    with open(input_file, 'r') as f:
        lines = f.readlines()

    for idx, line in enumerate(lines, 1):
        try:
            command, app_id, status = line.strip().split(',')

            app_url = f"{rm_base_url}/cluster/app/{app_id}"
            screenshot_path = os.path.join(screenshots_dir, f"{app_id}.png")

            # Open URL & take screenshot
            driver.get(app_url)
            time.sleep(2)  # Adjust if needed
            driver.save_screenshot(screenshot_path)

            # Add details and screenshot to Word doc
            doc.add_heading(f"Job {idx}", level=2)
            doc.add_paragraph(f"Command Executed: {command}")
            doc.add_paragraph(f"Status: {status}")
            doc.add_picture(screenshot_path, width=Inches(6))
            doc.add_paragraph("\n")

        except Exception as e:
            print(f"Error processing line {idx}: {line.strip()} - {e}")

    driver.quit()
    doc.save(output_doc)
    print(f"✅ Document saved: {output_doc}")


if __name__ == "__main__":
    # ✅ Update these values
    chrome_driver_path = r"C:\Users\YourUser\Downloads\chromedriver.exe"  # ← your actual path
    rm_base_url = "http://resourcemanager-host:8088"  # ← your ResourceManager base URL

    generate_doc("file1.txt", "success.docx", chrome_driver_path, rm_base_url)
    generate_doc("file2.txt", "failed.docx", chrome_driver_path, rm_base_url)
